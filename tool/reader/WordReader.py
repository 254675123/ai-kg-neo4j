# -*- coding: utf-8 -*-
# coding=utf-8
"""
create_author : zhangcl
create_time   : 2018-09-25
program       : *_* read word data  *_*

"""
import sys

import docx

import tool.processor.SentenceProcessor

# import xlwt
# from datetime import date,datetime
reload(sys)
sys.setdefaultencoding('utf-8')

class WordReader:
    """
    read data from word file.
    """

    def __init__(self):
        """
        initialize data
        """
        self.senpreprocessor = tool.processor.SentenceProcessor.SenPreprocess()
        self.input_filepath = None
        self.output_filepath = None
        self.docx_content = []
        self.result = []

    def readText(self):
        if self.input_filepath is None:
            return
        self.docx_content = []
        doc = docx.Document(self.input_filepath)
        index = 0
        for p in doc.paragraphs:
            index = index + 1
            one_row = p.text.strip()
            if len(one_row) == 0:
                continue


            # 不为空,存一份文本数据，做后期分析
            self.docx_content.append(one_row)

        # preprocess
        #self.preprocessResult(local_result)

        # 输出文本数据文件
        self.outputfile(self.output_filepath, self.docx_content)

    def readTextForVIP(self):
        if self.input_filepath is None:
            return
        local_result = []
        doc = docx.Document(self.input_filepath)
        index = 0
        for p in doc.paragraphs:
            index = index + 1
            one_row = p.text.strip()
            if len(one_row) == 0:
                continue

            # 不为空,存一份文本数据，做后期分析
            self.docx_content.append(one_row)

            if one_row.__contains__(u'经营决策要考虑的一些特定的成本概念有'):
                pass

            level,nline = self.senpreprocessor.judgeLevel(one_row)
            sens = []
            pbasestylename = '' if p.style.base_style is None else p.style.base_style.name
            pstylename = '' if p.style is None else p.style.name
            if (p._element.pPr is None or p._element.pPr.numPr is None):
                auto_numpr_level = -1
            else:
                auto_numpr_level = p._element.pPr.numPr.ilvl.val
            pstyleid = p.style.style_id
            custom_style = (pstyleid == pstylename and level > 0)
            if level > 0 and (pstylename.__contains__(u'Title') or pstylename.__contains__(u'Head') or pstylename.__contains__(u'标题') \
                    or pbasestylename.__contains__(u'Title') or pbasestylename.__contains__(u'Head')or pbasestylename.__contains__(u'标题')):
                local_result.append(one_row)
            elif level == 1 or level == 2:
                # level =1 说明是第一章之类的，level=2说明是第一节之类的
                # 对于目录之类的要处理页码
                ntext = self.senpreprocessor.removePageNum(one_row)
                local_result.append(ntext)
            elif custom_style:
                # 至少有一个是加粗的
                hasbold = self.hasBoldFont(p)
                if hasbold == False:
                    continue
                ntext = self.senpreprocessor.getPreSectionByIndex(one_row)
                if len(ntext) == 0:
                    continue
                local_result.append(ntext)
            elif auto_numpr_level >= 0:
                # 这属于自动编号的行，自动编号无法获取到
                # 自动编号的行，如果文本过长，超过30者，需要检查30个之内，是否有冒号，等号，或者破折号
                # 有的话，取前一部分，没有的话，说明是一段描述，不需要
                ntext = self.senpreprocessor.getPreSectionByIndex(one_row)
                if len(ntext) == 0:
                    continue
                auto_num = self.senpreprocessor.getAutoNumCode(auto_numpr_level)
                local_result.append(auto_num + ntext)

            else:
                for r in p.runs:
                    if len(r.text.strip()) == 0:
                        continue
                    # bold
                    if r.bold:
                        sens.append(r.text)
                        break

                if len(sens) > 0 and level > 0:
                    # 如果有句中有冒号，则只取冒号前的部分
                    line = self.scanSentence(one_row)
                    local_result.append(line)
                elif one_row.startswith(u'高校名称') or one_row.startswith(u'课程名称'):
                    local_result.append(one_row)

        # preprocess
        self.preprocessResult(local_result)

        # 输出文本数据文件
        self.outputfile(self.output_filepath, self.docx_content)
        self.docx_content = []


    def hasBoldFont(self,p):
        flag = False
        for r in p.runs:
            if len(r.text.strip()) == 0:
                continue
            # bold
            if r.bold:
                flag = True
                break
        return flag

    def preprocessResult(self, local_result):
        self.result = []
        for sen in local_result:

            sen_list = self.senpreprocessor.process(sen)
            if sen_list:
                self.result = self.result + sen_list
                #print sen_list[0]



    def scanSentence(self, sen):
        # 逐个字符扫描
        line_list = []
        for ch in sen:
            # 如果碰到冒号，则结束
            if ch == u':' or ch == u'：'or ch == u'。':
                break
            else:
                line_list.append(ch)

        return ''.join(line_list)

    def outputfile(self, filepath, content):
        if filepath is None:
            return
        fout = open(filepath, 'w')  # 以写得方式打开文件
        fout.write('\n'.join(content))  # 将分词好的结果写入到输出文件
        fout.close()

    def readTable(self):
        if self.input_filepath is None:
            return
        doc = docx.Document('tmp.docx')
        for table in doc.tables:  # 遍历所有表格
            print '----table------'
            for row in table.rows:  # 遍历表格的所有行
                # row_str = '\t'.join([cell.text for cell in row.cells])  # 一行数据
                # print row_str
                for cell in row.cells:
                    print cell.text, '\t',
                print

    def writeTable(self):

        doc = docx.Document()
        table = doc.add_table(rows=1, cols=3, style='Table Grid')  # 创建带边框的表格
        hdr_cells = table.rows[0].cells  # 获取第0行所有所有单元格
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'

        # 添加三行数据
        data_lines = 3
        for i in range(data_lines):
            cells = table.add_row().cells
            cells[0].text = 'Name%s' % i
            cells[1].text = 'Id%s' % i
            cells[2].text = 'Desc%s' % i

        rows = 2
        cols = 4
        table = doc.add_table(rows=rows, cols=cols)
        val = 1
        for i in range(rows):
            cells = table.rows[i].cells
            for j in range(cols):
                cells[j].text = str(val * 10)
                val += 1

        doc.save('tmp.docx')

if __name__ == "__main__":
    pusher = WordReader()
    pusher.input_filepath = u'D:/奥鹏/学生服务中心标注/文科课程电子辅导资料-docx/吉大《中级财务会计》.docx'
    pusher.readText()